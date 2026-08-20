from pathlib import Path

from app.domain.enums import DocumentType
from app.domain.schemas import DocumentRef, ParsedDocument
from app.parsing.interfaces import DocumentParser


class PilotFileParser(DocumentParser):
    """Local pilot parser used only when Docling is unavailable.

    It preserves the same DocumentParser contract and page provenance so the
    downstream verification path is exercised with the real PDF/DOCX bytes.
    Production remains configured to use DoclingAdapter.
    """

    def parse(self, document: DocumentRef) -> ParsedDocument:
        path = Path(document.storage_path)
        if not path.exists():
            raise ValueError(f"Pilot document not found: {path}")
        if path.suffix.lower() == ".pdf":
            from pypdf import PdfReader
            pages = {}
            for index, page in enumerate(PdfReader(str(path)).pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    raise ValueError(f"Pilot PDF page {index} contains no text")
                pages[index] = text
            return ParsedDocument(document_id=document.id, document_type=document.document_type, pages=pages)
        if path.suffix.lower() == ".docx":
            from docx import Document
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if not text.strip():
                raise ValueError("Pilot DOCX contains no text")
            return ParsedDocument(document_id=document.id, document_type=document.document_type, pages={1: text})
        raise ValueError(f"Unsupported pilot file type: {path.suffix}")
